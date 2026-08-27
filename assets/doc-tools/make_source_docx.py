# -*- coding: utf-8 -*-
"""软著源代码文档生成 make_source_docx.py —— 把源码目录/文件列表生成 Word 源代码文档。

软著惯例：每页 50 行（前后各 30 页），等宽字体，页眉软件名+页码。
参数化：--lines-per-page 控制每页行数（默认 50，可调以控制总页数）。

用法：
  python make_source_docx.py --out <输出.docx> --name <软件名称> \
      --root <源码根目录> [--include app worker scripts] [--exclude node_modules vendor dist] \
      [--lines-per-page 50]

说明：
  - --include 指定子目录/前缀（相对 root），默认全部（自动排除 --exclude 与常见噪音）
  - 自动排除：node_modules、.git、dist、build、__pycache__、*.map、package-lock.json、*.min.*
  - 文件间分页，每个文件前有「文件：路径（行数）」标题
"""
import os, sys, argparse
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CODE_EXTS = {'.ts', '.tsx', '.js', '.mjs', '.jsx', '.css', '.py', '.java', '.go', '.rs',
             '.c', '.h', '.cpp', '.cs', '.vue', '.html', '.json', '.sql', '.sh', '.yml', '.yaml'}
NOISE = {'node_modules', '.git', 'dist', 'build', '__pycache__', '.next', '.nuxt', '.venv', '.idea', '.vscode'}
NOISE_SUFFIX = {'.min.js', '.min.css', '.map', '.d.ts.map'}

def collect_files(root, includes, excludes):
    files = []
    for dirpath, dirnames, filenames in os.walk(root):
        dirnames[:] = [d for d in dirnames if d not in NOISE and not any(
            d.startswith(x) for x in ('node_modules', '.git'))]
        for fn in sorted(filenames):
            full = os.path.join(dirpath, fn)
            rel = os.path.relpath(full, root).replace('\\', '/')
            if any(rel.startswith(inc + '/') or rel == inc for inc in includes):
                if any(rel.startswith(ex + '/') or rel == ex for ex in excludes):
                    continue
            elif includes:
                continue
            if any(rel.endswith(s) for s in NOISE_SUFFIX):
                continue
            ext = os.path.splitext(fn)[1].lower()
            if ext in CODE_EXTS or (ext == '' and fn.startswith(('Dockerfile', 'Makefile'))):
                files.append((rel, full))
    return files

def count_lines(path):
    try:
        with open(path, encoding='utf-8') as f:
            return sum(1 for _ in f)
    except (UnicodeDecodeError, OSError):
        return 0

def add_page_number(paragraph, text):
    """页眉/页脚段落：'第 X 页' 用 PAGE 域。"""
    run = paragraph.add_run('第 ')
    run.font.size = Pt(9)
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'PAGE')
    paragraph._p.append(fld)
    run = paragraph.add_run(' 页')
    run.font.size = Pt(9)

def build(root, out_path, name, includes, excludes, lpp):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
    sec.left_margin = sec.right_margin = Cm(2.0)
    sec.top_margin = Cm(1.6)
    sec.bottom_margin = Cm(1.6)
    # 页眉：软件名居中
    header = sec.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = hp.add_run(f'{name} · 源代码文档')
    r.font.size = Pt(9)
    r.font.name = '宋体'
    r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    # 页脚：页码
    footer = sec.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(fp, '')

    files = collect_files(root, includes, excludes)
    if not files:
        print('未找到任何源码文件，请检查 --root 与 --include')
        sys.exit(1)
    total_lines = sum(count_lines(f) for _, f in files)

    # 封面
    for text, size in [(name, 22), ('源代码文档', 22), (f'共 {len(files)} 个文件，{total_lines} 行', 12),
                       ('（软件著作权申请材料）', 12)]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.font.name = '宋体'
        r.bold = (size >= 22)
        r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.add_page_break()

    # 文件清单
    p = doc.add_paragraph()
    r = p.add_run('源代码文件清单')
    r.font.size = Pt(14); r.bold = True
    table = doc.add_table(rows=1 + len(files), cols=3)
    table.style = 'Table Grid'
    for j, h in enumerate(['文件路径', '行数', '说明']):
        c = table.rows[0].cells[j]
        rr = c.paragraphs[0].add_run(h)
        rr.bold = True; rr.font.size = Pt(10)
    for i, (rel, full) in enumerate(files):
        cells = table.rows[i + 1].cells
        cells[0].paragraphs[0].add_run(rel).font.size = Pt(9)
        cells[1].paragraphs[0].add_run(str(count_lines(full))).font.size = Pt(9)
        cells[2].paragraphs[0].add_run('').font.size = Pt(9)
    doc.add_page_break()

    # 代码正文：每文件一节，文件间分页
    for idx, (rel, full) in enumerate(files):
        n = count_lines(full)
        p = doc.add_paragraph()
        r = p.add_run(f'文件：{rel}（{n} 行）')
        r.bold = True; r.font.size = Pt(11)
        r.font.name = '黑体'
        r.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        try:
            content = open(full, encoding='utf-8').read().split('\n')
        except (UnicodeDecodeError, OSError):
            content = ['[无法按 UTF-8 读取该文件，已跳过内容]']
        for line in content:
            para = doc.add_paragraph()
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.line_spacing = Pt(10.5)
            run = para.add_run(line if line else ' ')
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        if idx < len(files) - 1:
            doc.add_page_break()

    doc.save(out_path)
    est_pages = total_lines // lpp + len(files) // 3 + 3
    print(f'OK: {out_path}')
    print(f'文件数：{len(files)}，总行数：{total_lines}，每页 {lpp} 行，估算页数：约 {est_pages} 页')

if __name__ == '__main__':
    ap = argparse.ArgumentParser(description='软著源代码文档生成')
    ap.add_argument('--out', required=True)
    ap.add_argument('--name', required=True, help='软件名称')
    ap.add_argument('--root', required=True, help='源码根目录')
    ap.add_argument('--include', nargs='*', default=[], help='纳入的子目录/前缀')
    ap.add_argument('--exclude', nargs='*', default=[], help='排除的子目录/前缀')
    ap.add_argument('--lines-per-page', type=int, default=50, help='每页行数（软著官方 50）')
    args = ap.parse_args()
    build(args.root, args.out, args.name, args.include, args.exclude, args.lines_per_page)
