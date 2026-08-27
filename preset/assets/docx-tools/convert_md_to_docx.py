# -*- coding: utf-8 -*-
"""Markdown → Word 转换脚本（针对软著申报文档/技术文书的规整 md 结构）。
支持：标题层级 / 表格 / 代码块 / 图片嵌入（自动防溢出 A4）/ 有序无序列表 / 引用块 /
行内 **粗体** `代码` *斜体* / 图注居中加粗。

用法：修改下方 BASE 与 PAIRS（(md 文件, docx 输出名)），然后 python convert_md_to_docx.py
依赖：pip install python-docx pillow
"""
import re, os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from PIL import Image

# ── 配置 ────────────────────────────────────────────────────────────────────
BASE = r'.'  # md 文件所在目录（改成实际目录，如 D:\mini项目\docs）
PAIRS = [
    # ('文档名.md', '文档名.docx'),
]
MAX_W_CM = 14.5   # 横图限宽
MAX_H_CM = 22.0   # 竖长图限高（A4 页面高 29.7cm，留边距余量）
GRAY = RGBColor(0x59, 0x59, 0x59)

def fit_size(img_path):
    """按页面可用区域约束图片尺寸：横图限宽，竖长图限高（防止溢出 A4）。"""
    with Image.open(img_path) as im:
        w, h = im.size
    ratio = h / w
    if ratio > MAX_H_CM / MAX_W_CM:
        return Cm(MAX_H_CM / ratio), Cm(MAX_H_CM)
    return Cm(MAX_W_CM), Cm(MAX_W_CM * ratio)

def add_inline_runs(par, text):
    """解析行内 **bold**、`code`、*italic*，按片段添加 run。"""
    token_re = re.compile(r'(\*\*.+?\*\*|`[^`]+`|\*[^*]+?\*)')
    pos = 0
    for m in token_re.finditer(text):
        if m.start() > pos:
            par.add_run(text[pos:m.start()])
        tok = m.group(0)
        if tok.startswith('**'):
            r = par.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith('`'):
            r = par.add_run(tok[1:-1]); r.font.name = 'Consolas'; r.font.size = Pt(10)
        else:
            r = par.add_run(tok[1:-1]); r.italic = True
        pos = m.end()
    if pos < len(text):
        par.add_run(text[pos:])

def add_code_block(doc, lines):
    for i, line in enumerate(lines):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        r = p.add_run(line if line else ' ')
        r.font.name = 'Consolas'
        r.font.size = Pt(8)
        if i == 0:
            p.paragraph_format.space_before = Pt(4)
        if i == len(lines) - 1:
            p.paragraph_format.space_after = Pt(8)

def add_table(doc, rows):
    header = [c.strip() for c in rows[0].strip().strip('|').split('|')]
    body = []
    for row in rows[1:]:
        if re.match(r'^\s*\|?[\s:|-]+\|?\s*$', row):
            continue
        cells = [c.strip() for c in row.strip().strip('|').split('|')]
        body.append(cells)
    ncols = len(header)
    table = doc.add_table(rows=1 + len(body), cols=ncols)
    table.style = 'Table Grid'
    for j, cell_text in enumerate(header):
        cell = table.rows[0].cells[j]
        p = cell.paragraphs[0]
        add_inline_runs(p, cell_text)
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(10.5)
            if r.font.name is None:
                r.font.name = '宋体'
                r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    for i, row in enumerate(body):
        for j, val in enumerate(row):
            if j >= ncols:
                break
            cell = table.rows[i + 1].cells[j]
            p = cell.paragraphs[0]
            add_inline_runs(p, val)
            for r in p.runs:
                if r.font.size is None:
                    r.font.size = Pt(10.5)
    # 列宽均匀
    for j in range(ncols):
        for row in table.rows:
            row.cells[j].width = Cm(16.0 / ncols)

def convert(md_path, out_path):
    doc = Document()
    # 页面设置：A4、紧凑边距
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
    sec.left_margin = sec.right_margin = Cm(2.0)
    sec.top_margin = sec.bottom_margin = Cm(1.8)
    md_dir = os.path.dirname(os.path.abspath(md_path))
    pending_figcaption = False

    lines = open(md_path, encoding='utf-8').read().split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        s = line.strip()
        if not s or s == '---':
            i += 1; continue
        if s.startswith('```'):
            block = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                block.append(lines[i]); i += 1
            add_code_block(doc, block)
            i += 1; pending_figcaption = False; continue
        if s.startswith('|') and i + 1 < len(lines) and re.match(r'^\s*\|?[\s:|-]+\|?\s*$', lines[i + 1]):
            rows = [s]
            i += 1
            while i < len(lines) and lines[i].strip().startswith('|'):
                rows.append(lines[i].strip()); i += 1
            add_table(doc, rows)
            pending_figcaption = True; continue
        m = re.match(r'^(#{1,6})\s+(.*)$', s)
        if m:
            level = int(len(m.group(1)))
            p = doc.add_heading(level=min(level, 4))
            add_inline_runs(p, m.group(2))
            i += 1; pending_figcaption = False; continue
        m = re.match(r'^!\[(.*?)\]\((.*?)\)$', s)
        if m:
            alt, rel = m.group(1), m.group(2)
            img_path = rel if os.path.isabs(rel) else os.path.join(md_dir, rel)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if os.path.exists(img_path):
                run = p.add_run()
                w, h = fit_size(img_path)
                run.add_picture(img_path, width=w, height=h)
            else:
                p.add_run(f'[缺失图片: {rel}]')
            i += 1; pending_figcaption = True; continue
        # 图注/表注（紧跟图或表的单行「图N xxx」/「表N xxx」）→ 居中加粗
        if pending_figcaption and re.match(r'^(图|表)\d+', s):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(s)
            r.bold = True
            r.font.size = Pt(10.5)
            i += 1; pending_figcaption = False; continue
        if s.startswith('>'):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.75)
            r = p.add_run(s.lstrip('>').strip())
            r.italic = True
            r.font.color.rgb = GRAY
            i += 1; pending_figcaption = False; continue
        m = re.match(r'^[-*•]\s+(.*)$', s)
        if m:
            p = doc.add_paragraph(style='List Bullet')
            add_inline_runs(p, m.group(1))
            i += 1; pending_figcaption = False; continue
        m = re.match(r'^(\d+)[.、)]\s+(.*)$', s)
        if m:
            p = doc.add_paragraph(style='List Number')
            add_inline_runs(p, m.group(2))
            i += 1; pending_figcaption = False; continue
        p = doc.add_paragraph()
        add_inline_runs(p, s)
        i += 1; pending_figcaption = False
    doc.save(out_path)
    print(f'OK: {out_path}')

if __name__ == '__main__':
    for md, docx in PAIRS:
        convert(os.path.join(BASE, md), os.path.join(BASE, docx))
